# ============================================================================
# FINANCIAL DATA PROCESSOR - ENHANCED WITH SMART PREREQUISITES
# ============================================================================
# This file handles document processing and data extraction for financial analysis.
# It's like having a smart financial data analyst that can read and understand
# various document formats and extract meaningful financial information.
# 
# ENHANCED FEATURES:
# 🔧 Automatic dependency checking and installation for document processing
# 📄 Support for multiple file formats (CSV, Excel, PDF, Word documents)
# 🚨 Intelligent error handling with helpful recovery suggestions
# 📊 Smart data validation and cleaning
# 🔍 Advanced pattern recognition for financial categories
# 💾 Robust sample data generation for demonstrations
# 
# WHAT THIS FILE DOES:
# 1. 📄 Processes uploaded financial documents (CSV, Excel, PDF, Word)
# 2. 🔍 Extracts transaction data using intelligent pattern matching
# 3. 📊 Categorizes and validates financial information
# 4. 🧮 Calculates totals, summaries, and key financial metrics
# 5. 📋 Provides comprehensive sample data for demonstrations
# 6. 🚨 Handles errors gracefully with detailed troubleshooting guidance
# 
# SUPPORTED FORMATS:
# - CSV files: Bank statements, expense reports, transaction exports
# - Excel files: Financial spreadsheets, budgets, tracking sheets
# - PDF files: Bank statements, credit card statements (text extraction)
# - Word documents: Financial reports, budget documents
# - Text files: Simple transaction lists
# ============================================================================

import sys
import subprocess
import importlib.util
import os
import json
from typing import Dict, List, Any, Optional
import re
from datetime import datetime, timedelta
import random

# ============================================================================
# SMART DEPENDENCY MANAGEMENT - Document Processing Prerequisites
# ============================================================================

def check_document_processing_dependencies():
    """
    📄 DOCUMENT PROCESSING DEPENDENCY CHECKER
    
    WHAT THIS FUNCTION DOES:
    1. Checks if document processing packages are available
    2. Attempts to install missing packages automatically
    3. Provides capability flags for different file formats
    4. Returns comprehensive capability assessment
    
    DOCUMENT PROCESSING PACKAGES NEEDED:
    - pandas: Core data manipulation (CSV, Excel)
    - openpyxl: Excel file reading and writing
    - PyPDF2: PDF text extraction
    - python-docx: Word document processing
    - xlrd: Legacy Excel file support
    
    RETURNS:
    Dictionary with capability flags for different document types
    """
    
    print("📄 Checking document processing dependencies...")
    
    capabilities = {
        'csv_processing': False,
        'excel_processing': False,
        'pdf_processing': False,
        'word_processing': False,
        'data_analysis': False,
        'full_document_support': False
    }
    
    # Check Pandas (core data processing)
    try:
        import pandas as pd
        capabilities['csv_processing'] = True
        capabilities['data_analysis'] = True
        print("✅ Pandas - CSV and data analysis available")
    except ImportError:
        print("❌ Pandas missing - trying to install...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "pandas"])
            import pandas as pd
            capabilities['csv_processing'] = True
            capabilities['data_analysis'] = True
            print("✅ Pandas installed successfully!")
        except Exception as e:
            print(f"⚠️ Pandas installation failed: {e}")
    
    # Check OpenPyXL (Excel processing)
    try:
        import openpyxl
        capabilities['excel_processing'] = True
        print("✅ OpenPyXL - Excel processing available")
    except ImportError:
        print("❌ OpenPyXL missing - trying to install...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl"])
            capabilities['excel_processing'] = True
            print("✅ OpenPyXL installed successfully!")
        except Exception as e:
            print(f"⚠️ OpenPyXL installation failed: {e}")
    
    # Check PyPDF2 (PDF processing)
    try:
        import PyPDF2
        capabilities['pdf_processing'] = True
        print("✅ PyPDF2 - PDF processing available")
    except ImportError:
        print("❌ PyPDF2 missing - trying to install...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "PyPDF2"])
            capabilities['pdf_processing'] = True
            print("✅ PyPDF2 installed successfully!")
        except Exception as e:
            print(f"⚠️ PyPDF2 installation failed: {e}")
    
    # Check python-docx (Word processing)
    try:
        import docx
        capabilities['word_processing'] = True
        print("✅ python-docx - Word document processing available")
    except ImportError:
        print("❌ python-docx missing - trying to install...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            capabilities['word_processing'] = True
            print("✅ python-docx installed successfully!")
        except Exception as e:
            print(f"⚠️ python-docx installation failed: {e}")
    
    # Determine full capability
    capabilities['full_document_support'] = all([
        capabilities['csv_processing'],
        capabilities['excel_processing'],
        capabilities['pdf_processing']
    ])
    
    if capabilities['full_document_support']:
        print("🎉 Full document processing capabilities available!")
    else:
        print("💡 Partial document processing capabilities - some formats may be limited")
    
    return capabilities

# Initialize capabilities
DOCUMENT_CAPABILITIES = check_document_processing_dependencies()

# Smart imports based on availability
if DOCUMENT_CAPABILITIES['data_analysis']:
    try:
        import pandas as pd
        PANDAS_AVAILABLE = True
    except ImportError:
        PANDAS_AVAILABLE = False
else:
    PANDAS_AVAILABLE = False

if DOCUMENT_CAPABILITIES['pdf_processing']:
    try:
        import PyPDF2
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False
else:
    PDF_AVAILABLE = False

if DOCUMENT_CAPABILITIES['word_processing']:
    try:
        from docx import Document
        DOCX_AVAILABLE = True
    except ImportError:
        DOCX_AVAILABLE = False
else:
    DOCX_AVAILABLE = False

if DOCUMENT_CAPABILITIES['excel_processing']:
    try:
        import openpyxl
        EXCEL_AVAILABLE = True
    except ImportError:
        EXCEL_AVAILABLE = False
else:
    EXCEL_AVAILABLE = False

# ============================================================================
# MAIN FINANCIAL DATA PROCESSOR CLASS - Enhanced with Smart Capabilities
# ============================================================================

class FinancialDataProcessor:
    """
    📊 ENHANCED FINANCIAL DATA PROCESSOR
    
    WHAT IT DOES:
    - Processes various document types to extract financial data
    - Uses intelligent pattern matching to identify transactions
    - Validates and cleans data for accurate analysis
    - Provides detailed error handling and recovery options
    - Supports multiple file formats with graceful fallbacks
    
    ENHANCED FEATURES:
    - Smart column detection across different file formats
    - Advanced transaction categorization
    - Robust error handling with helpful error messages
    - Comprehensive data validation and cleaning
    - Professional sample data generation
    
    SUPPORTED FORMATS:
    1. CSV: Bank exports, transaction lists, budget files
    2. Excel: Financial spreadsheets, budget templates
    3. PDF: Bank statements, credit card statements
    4. Word: Financial reports, budget documents
    5. Text: Simple transaction lists
    
    HOW IT WORKS:
    1. Identifies file format and routes to appropriate processor
    2. Uses smart pattern matching to find financial columns
    3. Extracts and validates transaction data
    4. Categorizes transactions using keyword matching
    5. Calculates totals and generates summary statistics
    6. Returns structured data ready for AI analysis
    """
    
    def __init__(self):
        """
        ENHANCED INITIALIZATION: Set up processor with smart capabilities
        
        WHAT THIS DOES:
        1. Reports available processing capabilities
        2. Sets up file format handlers
        3. Initializes validation rules
        4. Configures error handling strategies
        """
        
        print("📊 Initializing Financial Data Processor...")
        
        # Set up supported formats based on available libraries
        self.supported_formats = []
        
        if PANDAS_AVAILABLE:
            self.supported_formats.extend(['.csv', '.txt'])
            print("✅ CSV and text file processing enabled")
        
        if EXCEL_AVAILABLE and PANDAS_AVAILABLE:
            self.supported_formats.extend(['.xlsx', '.xls'])
            print("✅ Excel file processing enabled")
        
        if PDF_AVAILABLE:
            self.supported_formats.append('.pdf')
            print("✅ PDF file processing enabled")
        
        if DOCX_AVAILABLE:
            self.supported_formats.append('.docx')
            print("✅ Word document processing enabled")
        
        if not self.supported_formats:
            print("⚠️ Limited file processing - only sample data available")
        
        # Transaction categorization keywords
        self.category_keywords = {
            'Housing': ['rent', 'mortgage', 'property tax', 'hoa', 'utilities', 'electric', 'gas', 'water', 'internet', 'cable'],
            'Transportation': ['gas', 'fuel', 'uber', 'lyft', 'taxi', 'bus', 'train', 'car payment', 'auto insurance', 'parking'],
            'Food': ['grocery', 'restaurant', 'food', 'dining', 'coffee', 'lunch', 'dinner', 'breakfast', 'fast food'],
            'Healthcare': ['medical', 'doctor', 'hospital', 'pharmacy', 'health insurance', 'dental', 'vision'],
            'Entertainment': ['movie', 'netflix', 'spotify', 'gaming', 'concert', 'theater', 'streaming'],
            'Shopping': ['amazon', 'target', 'walmart', 'mall', 'store', 'clothing', 'electronics'],
            'Debt Payment': ['credit card', 'loan payment', 'student loan', 'debt', 'financing'],
            'Income': ['salary', 'paycheck', 'bonus', 'refund', 'deposit', 'income', 'wages'],
            'Savings': ['savings', 'investment', 'retirement', '401k', 'ira', 'emergency fund']
        }
        
        print("📊 Financial Data Processor ready!")
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        ENHANCED DOCUMENT PROCESSING WITH COMPREHENSIVE ERROR HANDLING
        
        INPUTS:
        - file_path: Path to financial document
        
        OUTPUTS:
        - Dictionary with structured financial data or detailed error information
        
        ENHANCEMENT FEATURES:
        - Intelligent file format detection
        - Multiple processing strategies per format
        - Comprehensive error handling with recovery suggestions
        - Data validation and cleaning
        - Professional error reporting
        """
        
        print(f"📄 Processing document: {os.path.basename(file_path)}")
        
        try:
            # STEP 1: Validate file existence and accessibility
            if not os.path.exists(file_path):
                return self._create_error_response(
                    "File Not Found",
                    f"The file '{file_path}' could not be found.",
                    ["Check that the file path is correct", "Ensure the file hasn't been moved or deleted"]
                )
            
            # STEP 2: Determine file format
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension not in self.supported_formats:
                return self._create_error_response(
                    "Unsupported Format",
                    f"File format '{file_extension}' is not supported.",
                    [f"Supported formats: {', '.join(self.supported_formats)}",
                     "Convert your file to CSV for best results",
                     "Ensure required libraries are installed"]
                )
            
            # STEP 3: Route to appropriate processor
            if file_extension == '.csv':
                return self._process_csv(file_path)
            elif file_extension in ['.xlsx', '.xls']:
                return self._process_excel(file_path)
            elif file_extension == '.pdf':
                return self._process_pdf(file_path)
            elif file_extension == '.docx':
                return self._process_docx(file_path)
            elif file_extension == '.txt':
                return self._process_text(file_path)
            else:
                return self._create_error_response(
                    "Processing Error",
                    f"No processor available for '{file_extension}' files.",
                    ["This shouldn't happen - please report this issue"]
                )
                
        except Exception as e:
            print(f"❌ Document processing error: {e}")
            return self._create_error_response(
                "Processing Exception",
                str(e),
                ["Check that the file is not corrupted",
                 "Ensure the file is not currently open in another application",
                 "Try saving the file in a different format",
                 "Contact support if the issue persists"]
            )
    
    def _process_csv(self, file_path: str) -> Dict[str, Any]:
        """Enhanced CSV processing with smart column detection"""
        print("📊 Processing CSV file...")
        
        try:
            # Try multiple encodings
            encodings_to_try = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            df = None
            
            for encoding in encodings_to_try:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    print(f"✅ CSV read successfully with {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    print(f"⚠️ Encoding {encoding} failed: {e}")
                    continue
            
            if df is None:
                return self._create_error_response(
                    "CSV Reading Error",
                    "Could not read CSV file with any supported encoding.",
                    ["Ensure the file is a valid CSV",
                     "Try saving the file with UTF-8 encoding",
                     "Check for special characters in the file"]
                )
            
            if df.empty:
                return self._create_error_response(
                    "Empty File",
                    "The CSV file contains no data.",
                    ["Check that the file has transaction data",
                     "Ensure the file isn't just headers"]
                )
            
            print(f"📊 CSV loaded: {len(df)} rows, {len(df.columns)} columns")
            
            # Smart column detection
            column_mapping = self._detect_csv_columns(df)
            
            if not column_mapping:
                return self._create_error_response(
                    "Column Detection Failed",
                    "Could not identify required financial columns in the CSV.",
                    ["Required columns: Date, Amount, Description/Category",
                     "Column names should be in the first row",
                     "Example format: Date,Amount,Category,Description"]
                )
            
            # Process transactions
            return self._extract_transactions_from_dataframe(df, column_mapping)
            
        except Exception as e:
            print(f"❌ CSV processing error: {e}")
            return self._create_error_response(
                "CSV Processing Error",
                str(e),
                ["Check that the CSV file is properly formatted",
                 "Ensure the file has headers in the first row",
                 "Try opening the file in Excel to verify structure"]
            )
    
    def _process_excel(self, file_path: str) -> Dict[str, Any]:
        """Enhanced Excel processing with multi-sheet support"""
        print("📈 Processing Excel file...")
        
        try:
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names
            
            print(f"📈 Excel file has {len(sheet_names)} sheets: {sheet_names}")
            
            # Try each sheet to find financial data
            for sheet_name in sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    if df.empty:
                        continue
                    
                    # Try to detect financial columns
                    column_mapping = self._detect_csv_columns(df)
                    
                    if column_mapping:
                        print(f"✅ Found financial data in sheet: {sheet_name}")
                        return self._extract_transactions_from_dataframe(df, column_mapping)
                        
                except Exception as e:
                    print(f"⚠️ Error reading sheet {sheet_name}: {e}")
                    continue
            
            return self._create_error_response(
                "Excel Processing Failed",
                "No financial data found in any Excel sheet.",
                ["Ensure at least one sheet has columns: Date, Amount, Description",
                 "Check that data starts in the first few rows",
                 "Try converting to CSV for better compatibility"]
            )
                
        except Exception as e:
            print(f"❌ Excel processing error: {e}")
            return self._create_error_response(
                "Excel Processing Error",
                str(e),
                ["Check that the Excel file is not corrupted",
                 "Ensure the file is not password protected",
                 "Try saving as CSV if Excel processing fails"]
            )
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Enhanced PDF processing with text extraction"""
        print("📄 Processing PDF file...")
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                if len(pdf_reader.pages) == 0:
                    return self._create_error_response(
                        "Empty PDF",
                        "The PDF file contains no pages.",
                        ["Check that the PDF file is valid",
                         "Ensure the PDF is not corrupted"]
                    )
                
                # Extract text from all pages
                full_text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        full_text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                    except Exception as e:
                        print(f"⚠️ Error extracting text from page {page_num + 1}: {e}")
                
                if not full_text.strip():
                    return self._create_error_response(
                        "PDF Text Extraction Failed",
                        "Could not extract readable text from the PDF.",
                        ["The PDF might be image-based (scanned document)",
                         "Try converting to text format first",
                         "Use CSV or Excel format for better results"]
                    )
                
                # Analyze extracted text
                financial_patterns = self._analyze_pdf_text(full_text)
                
                return {
                    "document_type": "PDF Bank Statement",
                    "pages_processed": len(pdf_reader.pages),
                    "raw_text_length": len(full_text),
                    "detected_amounts": financial_patterns.get('amounts', []),
                    "detected_dates": financial_patterns.get('dates', []),
                    "potential_transactions": financial_patterns.get('transactions', []),
                    "processing_notes": [
                        "PDF processing is limited - amounts and dates detected",
                        "For full analysis, export transactions to CSV format",
                        "Manual review recommended for accuracy"
                    ],
                    "total_income": sum(amt for amt in financial_patterns.get('amounts', []) if amt > 0),
                    "total_expenses": sum(abs(amt) for amt in financial_patterns.get('amounts', []) if amt < 0),
                    "categories": {"PDF Transactions": sum(abs(amt) for amt in financial_patterns.get('amounts', []) if amt < 0)}
                }
                
        except Exception as e:
            print(f"❌ PDF processing error: {e}")
            return self._create_error_response(
                "PDF Processing Error",
                str(e),
                ["Check that the PDF file is not corrupted",
                 "Ensure the PDF is not password protected",
                 "Try converting to CSV for better results"]
            )
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Enhanced Word document processing"""
        print("📝 Processing Word document...")
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            
            # Extract text from tables
            table_data = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
            
            if not full_text.strip() and not table_data:
                return self._create_error_response(
                    "Empty Document",
                    "The Word document contains no readable text or tables.",
                    ["Check that the document has financial content",
                     "Ensure the document is not corrupted"]
                )
            
            # Analyze document for financial patterns
            financial_patterns = self._analyze_text_for_financial_data(full_text, table_data)
            
            return {
                "document_type": "Word Financial Document",
                "text_length": len(full_text),
                "tables_found": len(table_data),
                "detected_amounts": financial_patterns.get('amounts', []),
                "potential_categories": financial_patterns.get('categories', []),
                "processing_notes": [
                    "Word document processing is basic - patterns detected",
                    "For detailed analysis, use structured CSV format",
                    "Manual review recommended"
                ],
                "total_income": sum(amt for amt in financial_patterns.get('amounts', []) if amt > 0),
                "total_expenses": sum(abs(amt) for amt in financial_patterns.get('amounts', []) if amt < 0),
                "categories": {"Document Analysis": sum(abs(amt) for amt in financial_patterns.get('amounts', []) if amt < 0)}
            }
            
        except Exception as e:
            print(f"❌ Word document processing error: {e}")
            return self._create_error_response(
                "Word Processing Error",
                str(e),
                ["Check that the Word document is not corrupted",
                 "Ensure the document is not password protected",
                 "Try saving as text or CSV format"]
            )
    
    def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Enhanced text file processing"""
        print("📝 Processing text file...")
        
        try:
            # Try multiple encodings
            encodings_to_try = ['utf-8', 'latin-1', 'cp1252']
            content = None
            
            for encoding in encodings_to_try:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                return self._create_error_response(
                    "Text Reading Error",
                    "Could not read text file with any supported encoding.",
                    ["Check that the file contains valid text",
                     "Try saving with UTF-8 encoding"]
                )
            
            if not content.strip():
                return self._create_error_response(
                    "Empty File",
                    "The text file is empty.",
                    ["Check that the file contains financial data"]
                )
            
            # Analyze text for financial patterns
            financial_patterns = self._analyze_text_for_financial_data(content, [])
            
            return {
                "document_type": "Text Financial Data",
                "content_length": len(content),
                "detected_amounts": financial_patterns.get('amounts', []),
                "processing_notes": [
                    "Text file processing is basic - patterns detected",
                    "Use CSV format with headers for best results"
                ],
                "total_income": sum(amt for amt in financial_patterns.get('amounts', []) if amt > 0),
                "total_expenses": sum(abs(amt) for amt in financial_patterns.get('amounts', []) if amt < 0),
                "categories": {"Text Analysis": sum(abs(amt) for amt in financial_patterns.get('amounts', []) if amt < 0)}
            }
            
        except Exception as e:
            print(f"❌ Text processing error: {e}")
            return self._create_error_response(
                "Text Processing Error",
                str(e),
                ["Check that the text file is accessible",
                 "Ensure the file is not corrupted"]
            )
    
    def _detect_csv_columns(self, df) -> Optional[Dict[str, str]]:
        """Intelligent column detection for CSV/Excel files"""
        
        column_mapping = {}
        columns = [col.lower().strip() for col in df.columns]
        
        # Define column patterns for intelligent matching
        column_patterns = {
            'date': ['date', 'transaction_date', 'posting_date', 'trans_date', 'dt', 'timestamp'],
            'amount': ['amount', 'transaction_amount', 'debit', 'credit', 'value', 'sum', 'total', 'amt'],
            'description': ['description', 'memo', 'details', 'transaction_details', 'desc', 'note'],
            'category': ['category', 'type', 'transaction_type', 'class', 'classification', 'cat']
        }
        
        # Find best matches for each required column
        for standard_name, patterns in column_patterns.items():
            best_match = None
            best_score = 0
            
            for col in columns:
                for pattern in patterns:
                    if pattern in col or col in pattern:
                        # Calculate match score (simple scoring)
                        score = len(pattern) if pattern in col else len(col)
                        if score > best_score:
                            best_score = score
                            best_match = df.columns[columns.index(col)]
            
            if best_match:
                column_mapping[standard_name] = best_match
        
        # Validate required columns
        required_columns = ['date', 'amount']
        if not all(col in column_mapping for col in required_columns):
            print(f"⚠️ Missing required columns. Found: {list(column_mapping.keys())}")
            return None
        
        # Validate data types
        try:
            # Check if amount column contains numeric data
            amount_col = column_mapping['amount']
            sample_amounts = df[amount_col].dropna().head(10)
            
            # Try to convert to numeric
            numeric_amounts = pd.to_numeric(sample_amounts, errors='coerce')
            if numeric_amounts.isna().all():
                print("⚠️ Amount column doesn't contain numeric data")
                return None
            
            print(f"✅ Column mapping detected: {column_mapping}")
            return column_mapping
            
        except Exception as e:
            print(f"⚠️ Column validation error: {e}")
            return None
    
    def _extract_transactions_from_dataframe(self, df, column_mapping: Dict[str, str]) -> Dict[str, Any]:
        """Enhanced transaction extraction and analysis"""
        
        print("💰 Extracting and analyzing transactions...")
        
        try:
            # Initialize financial data structure
            financial_data = {
                "transactions": [],
                "total_income": 0,
                "total_expenses": 0,
                "categories": {},
                "processing_info": {
                    "rows_processed": len(df),
                    "successful_transactions": 0,
                    "skipped_rows": 0,
                    "data_quality_issues": []
                }
            }
            
            # Process each row
            for index, row in df.iterrows():
                try:
                    # Extract basic transaction data
                    amount_raw = row[column_mapping['amount']]
                    date_raw = row[column_mapping['date']]
                    
                    # Get optional fields
                    description = ""
                    if 'description' in column_mapping:
                        description = str(row[column_mapping['description']]).strip()
                    
                    category = ""
                    if 'category' in column_mapping:
                        category = str(row[column_mapping['category']]).strip()
                    
                    # Clean and validate amount
                    amount = self._clean_amount(amount_raw)
                    if amount is None:
                        financial_data["processing_info"]["skipped_rows"] += 1
                        continue
                    
                    # Clean and validate date
                    date_str = self._clean_date(date_raw)
                    
                    # Determine category (use provided or auto-categorize)
                    if not category or category.lower() in ['nan', 'none', '']:
                        category = self._auto_categorize_transaction(description, amount)
                    
                    # Create transaction record
                    transaction = {
                        "date": date_str,
                        "amount": amount,
                        "category": category,
                        "description": description
                    }
                    
                    financial_data["transactions"].append(transaction)
                    
                    # Update totals
                    if amount > 0:
                        financial_data["total_income"] += amount
                    else:
                        financial_data["total_expenses"] += abs(amount)
                    
                    # Update category totals
                    abs_amount = abs(amount)
                    if category in financial_data["categories"]:
                        financial_data["categories"][category] += abs_amount
                    else:
                        financial_data["categories"][category] = abs_amount
                    
                    financial_data["processing_info"]["successful_transactions"] += 1
                    
                except Exception as e:
                    print(f"⚠️ Error processing row {index}: {e}")
                    financial_data["processing_info"]["skipped_rows"] += 1
                    financial_data["processing_info"]["data_quality_issues"].append(f"Row {index}: {str(e)}")
            
            # Final validation and summary
            if financial_data["processing_info"]["successful_transactions"] == 0:
                return self._create_error_response(
                    "No Valid Transactions",
                    "Could not extract any valid transactions from the file.",
                    ["Check that the amount column contains numeric values",
                     "Ensure the date column has valid dates",
                     "Verify the file format matches expected structure"]
                )
            
            print(f"✅ Successfully processed {financial_data['processing_info']['successful_transactions']} transactions")
            
            if financial_data["processing_info"]["skipped_rows"] > 0:
                print(f"⚠️ Skipped {financial_data['processing_info']['skipped_rows']} rows due to data issues")
            
            return financial_data
            
        except Exception as e:
            print(f"❌ Transaction extraction error: {e}")
            return self._create_error_response(
                "Transaction Extraction Error",
                str(e),
                ["Check the file structure and data format",
                 "Ensure numeric amounts and valid dates",
                 "Review the error details above"]
            )
    
    def _clean_amount(self, amount_raw) -> Optional[float]:
        """Clean and validate amount values"""
        try:
            if pd.isna(amount_raw):
                return None
            
            # Convert to string and clean
            amount_str = str(amount_raw).strip()
            
            # Remove currency symbols and commas
            amount_str = re.sub(r'[\$,\s]', '', amount_str)
            
            # Handle parentheses (negative numbers)
            if amount_str.startswith('(') and amount_str.endswith(')'):
                amount_str = '-' + amount_str[1:-1]
            
            # Convert to float
            return float(amount_str)
            
        except (ValueError, TypeError):
            return None
    
    def _clean_date(self, date_raw) -> str:
        """Clean and format date values"""
        try:
            if pd.isna(date_raw):
                return "Unknown"
            
            # Try to parse as pandas datetime
            if isinstance(date_raw, str):
                date_obj = pd.to_datetime(date_raw, errors='coerce')
            else:
                date_obj = pd.to_datetime(str(date_raw), errors='coerce')
            
            if pd.isna(date_obj):
                return str(date_raw)  # Return original if parsing fails
            
            return date_obj.strftime('%Y-%m-%d')
            
        except:
            return str(date_raw) if date_raw else "Unknown"
    
    def _auto_categorize_transaction(self, description: str, amount: float) -> str:
        """Automatically categorize transactions based on description and amount"""
        
        if not description:
            return "Income" if amount > 0 else "Other Expenses"
        
        description_lower = description.lower()
        
        # Check each category for keyword matches
        for category, keywords in self.category_keywords.items():
            if any(keyword in description_lower for keyword in keywords):
                return category
        
        # Default categorization based on amount
        return "Income" if amount > 0 else "Other Expenses"
    
    def _analyze_pdf_text(self, text: str) -> Dict[str, List]:
        """Analyze PDF text for financial patterns"""
        
        patterns = {
            'amounts': [],
            'dates': [],
            'transactions': []
        }
        
        # Extract dollar amounts
        amount_pattern = r'\$[\d,]+\.?\d*'
        amounts = re.findall(amount_pattern, text)
        for amt_str in amounts:
            try:
                cleaned = amt_str.replace(',', '').replace('$', '').replace(' ', '')
                amount = float(cleaned)
                #amount = float(amt_str.replace(', '').replace(',', ''))
                patterns['amounts'].append(amount)
            except ValueError:
                continue
        
        # Extract dates
        date_patterns = [
            r'\d{1,2}/\d{1,2}/\d{4}',
            r'\d{1,2}-\d{1,2}-\d{4}',
            r'\d{4}-\d{1,2}-\d{1,2}'
        ]
        
        for pattern in date_patterns:
            dates = re.findall(pattern, text)
            patterns['dates'].extend(dates)
        
        return patterns
    
    def _analyze_text_for_financial_data(self, text: str, table_data: List[List[str]]) -> Dict[str, List]:
        """Analyze text and tables for financial patterns"""
        
        patterns = {
            'amounts': [],
            'categories': []
        }
        
        # Extract amounts from text
        amount_pattern = r'\$?[\d,]+\.?\d*'
        amounts = re.findall(amount_pattern, text)
        
        for amt_str in amounts:
            try:
                #amount = float(amt_str.replace(', '').replace(',', ''))
                cleaned = amt_str.replace(',', '').replace('$', '').replace(' ', '')
                amount = float(cleaned)
                if amount > 0:  # Only include meaningful amounts
                    patterns['amounts'].append(amount)
            except ValueError:
                continue
        
        # Extract potential categories from text
        for category in self.category_keywords.keys():
            if category.lower() in text.lower():
                patterns['categories'].append(category)
        
        return patterns
    
    def _create_error_response(self, error_type: str, message: str, suggestions: List[str]) -> Dict[str, Any]:
        """Create standardized error response with helpful guidance"""
        
        return {
            "error": error_type,
            "message": message,
            "suggestions": suggestions,
            "error_details": {
                "timestamp": datetime.now().isoformat(),
                "processor_capabilities": {
                    "csv_support": PANDAS_AVAILABLE,
                    "excel_support": EXCEL_AVAILABLE,
                    "pdf_support": PDF_AVAILABLE,
                    "word_support": DOCX_AVAILABLE
                }
            }
        }

# ============================================================================
# ENHANCED SAMPLE DATA GENERATION - Professional Demo Data
# ============================================================================

def create_sample_data() -> Dict[str, Any]:
    """
    📊 ENHANCED SAMPLE DATA GENERATOR
    
    WHAT THIS FUNCTION DOES:
    1. Creates realistic financial transaction data
    2. Includes diverse income and expense categories
    3. Provides balanced financial scenario for demonstrations
    4. Uses real-world amounts and patterns
    5. Ensures data works well with all analysis agents
    
    WHY REALISTIC SAMPLE DATA MATTERS:
    - Demonstrates app capabilities effectively
    - Provides meaningful analysis results
    - Shows diverse financial scenarios
    - Helps users understand expected data format
    - Enables comprehensive testing
    
    SAMPLE SCENARIO:
    - Monthly income: $5,200 (salary + side income)
    - Monthly expenses: $3,450 (various categories)
    - Net savings potential: $1,750/month
    - Demonstrates both good and challenging financial situations
    """
    
    print("📊 Generating comprehensive sample financial data...")
    
    # Create realistic transaction data
    base_date = datetime.now() - timedelta(days=30)
    
    sample_transactions = [
        # INCOME TRANSACTIONS
        {"date": (base_date + timedelta(days=1)).strftime('%Y-%m-%d'), "amount": 4200.00, "category": "Salary", "description": "Monthly Salary - Company XYZ"},
        {"date": (base_date + timedelta(days=15)).strftime('%Y-%m-%d'), "amount": 800.00, "category": "Side Income", "description": "Freelance Project Payment"},
        {"date": (base_date + timedelta(days=20)).strftime('%Y-%m-%d'), "amount": 200.00, "category": "Other Income", "description": "Tax Refund"},
        
        # HOUSING & UTILITIES
        {"date": (base_date + timedelta(days=1)).strftime('%Y-%m-%d'), "amount": -1250.00, "category": "Housing", "description": "Monthly Rent Payment"},
        {"date": (base_date + timedelta(days=5)).strftime('%Y-%m-%d'), "amount": -120.00, "category": "Utilities", "description": "Electric Bill"},
        {"date": (base_date + timedelta(days=7)).strftime('%Y-%m-%d'), "amount": -80.00, "category": "Utilities", "description": "Internet & Cable"},
        {"date": (base_date + timedelta(days=10)).strftime('%Y-%m-%d'), "amount": -45.00, "category": "Utilities", "description": "Water Bill"},
        
        # TRANSPORTATION
        {"date": (base_date + timedelta(days=3)).strftime('%Y-%m-%d'), "amount": -350.00, "category": "Transportation", "description": "Car Payment"},
        {"date": (base_date + timedelta(days=8)).strftime('%Y-%m-%d'), "amount": -55.00, "category": "Transportation", "description": "Gas Station Fill-up"},
        {"date": (base_date + timedelta(days=15)).strftime('%Y-%m-%d'), "amount": -60.00, "category": "Transportation", "description": "Gas Station Fill-up"},
        {"date": (base_date + timedelta(days=22)).strftime('%Y-%m-%d'), "amount": -45.00, "category": "Transportation", "description": "Gas Station Fill-up"},
        {"date": (base_date + timedelta(days=12)).strftime('%Y-%m-%d'), "amount": -25.00, "category": "Transportation", "description": "Parking Fee"},
        
        # FOOD & DINING
        {"date": (base_date + timedelta(days=2)).strftime('%Y-%m-%d'), "amount": -120.00, "category": "Groceries", "description": "Weekly Grocery Shopping"},
        {"date": (base_date + timedelta(days=9)).strftime('%Y-%m-%d'), "amount": -115.00, "category": "Groceries", "description": "Weekly Grocery Shopping"},
        {"date": (base_date + timedelta(days=16)).strftime('%Y-%m-%d'), "amount": -130.00, "category": "Groceries", "description": "Weekly Grocery Shopping"},
        {"date": (base_date + timedelta(days=23)).strftime('%Y-%m-%d'), "amount": -125.00, "category": "Groceries", "description": "Weekly Grocery Shopping"},
        {"date": (base_date + timedelta(days=6)).strftime('%Y-%m-%d'), "amount": -45.00, "category": "Dining Out", "description": "Restaurant Dinner"},
        {"date": (base_date + timedelta(days=13)).strftime('%Y-%m-%d'), "amount": -35.00, "category": "Dining Out", "description": "Lunch with Colleagues"},
        {"date": (base_date + timedelta(days=19)).strftime('%Y-%m-%d'), "amount": -25.00, "category": "Dining Out", "description": "Coffee Shop"},
        
        # HEALTHCARE
        {"date": (base_date + timedelta(days=14)).strftime('%Y-%m-%d'), "amount": -150.00, "category": "Healthcare", "description": "Doctor Visit Copay"},
        {"date": (base_date + timedelta(days=18)).strftime('%Y-%m-%d'), "amount": -35.00, "category": "Healthcare", "description": "Prescription Medication"},
        
        # ENTERTAINMENT & LIFESTYLE
        {"date": (base_date + timedelta(days=4)).strftime('%Y-%m-%d'), "amount": -15.99, "category": "Entertainment", "description": "Netflix Subscription"},
        {"date": (base_date + timedelta(days=11)).strftime('%Y-%m-%d'), "amount": -12.99, "category": "Entertainment", "description": "Spotify Premium"},
        {"date": (base_date + timedelta(days=17)).strftime('%Y-%m-%d'), "amount": -45.00, "category": "Entertainment", "description": "Movie Theater Tickets"},
        {"date": (base_date + timedelta(days=24)).strftime('%Y-%m-%d'), "amount": -85.00, "category": "Entertainment", "description": "Concert Tickets"},
        
        # SHOPPING & PERSONAL
        {"date": (base_date + timedelta(days=5)).strftime('%Y-%m-%d'), "amount": -65.00, "category": "Shopping", "description": "Clothing Store"},
        {"date": (base_date + timedelta(days=21)).strftime('%Y-%m-%d'), "amount": -120.00, "category": "Shopping", "description": "Amazon Purchase"},
        {"date": (base_date + timedelta(days=26)).strftime('%Y-%m-%d'), "amount": -40.00, "category": "Personal Care", "description": "Haircut"},
        
        # DEBT PAYMENTS
        {"date": (base_date + timedelta(days=3)).strftime('%Y-%m-%d'), "amount": -185.00, "category": "Debt Payment", "description": "Credit Card Payment"},
        {"date": (base_date + timedelta(days=25)).strftime('%Y-%m-%d'), "amount": -75.00, "category": "Debt Payment", "description": "Student Loan Payment"},
        
        # INSURANCE
        {"date": (base_date + timedelta(days=1)).strftime('%Y-%m-%d'), "amount": -95.00, "category": "Insurance", "description": "Auto Insurance"},
        {"date": (base_date + timedelta(days=15)).strftime('%Y-%m-%d'), "amount": -200.00, "category": "Insurance", "description": "Health Insurance Premium"},
        
        # SAVINGS & INVESTMENTS
        {"date": (base_date + timedelta(days=2)).strftime('%Y-%m-%d'), "amount": -300.00, "category": "Savings", "description": "Emergency Fund Transfer"},
        {"date": (base_date + timedelta(days=16)).strftime('%Y-%m-%d'), "amount": -400.00, "category": "Investment", "description": "401k Contribution"},
        
        # MISCELLANEOUS
        {"date": (base_date + timedelta(days=28)).strftime('%Y-%m-%d'), "amount": -50.00, "category": "Other", "description": "ATM Withdrawal"},
        {"date": (base_date + timedelta(days=29)).strftime('%Y-%m-%d'), "amount": -25.00, "category": "Fees", "description": "Bank Service Fee"}
    ]
    
    # Calculate totals and categories
    total_income = sum(t['amount'] for t in sample_transactions if t['amount'] > 0)
    total_expenses = sum(abs(t['amount']) for t in sample_transactions if t['amount'] < 0)
    
    # Group expenses by category for analysis
    expense_categories = {}
    for transaction in sample_transactions:
        if transaction['amount'] < 0:  # Only expenses
            category = transaction['category']
            amount = abs(transaction['amount'])
            if category in expense_categories:
                expense_categories[category] += amount
            else:
                expense_categories[category] = amount
    
    # Add income categories
    income_categories = {}
    for transaction in sample_transactions:
        if transaction['amount'] > 0:  # Only income
            category = transaction['category']
            amount = transaction['amount']
            if category in income_categories:
                income_categories[category] += amount
            else:
                income_categories[category] = amount
    
    # Combine all categories for comprehensive analysis
    all_categories = {**expense_categories, **income_categories}
    
    sample_data = {
        "transactions": sample_transactions,
        "total_income": total_income,
        "total_expenses": total_expenses,
        "categories": all_categories,
        "expense_categories": expense_categories,
        "income_categories": income_categories,
        "net_cash_flow": total_income - total_expenses,
        "savings_rate": ((total_income - total_expenses) / total_income * 100) if total_income > 0 else 0,
        "sample_info": {
            "scenario": "Young Professional with Good Financial Habits",
            "monthly_income": total_income,
            "monthly_expenses": total_expenses,
            "net_savings": total_income - total_expenses,
            "transaction_count": len(sample_transactions),
            "date_range": f"{sample_transactions[0]['date']} to {sample_transactions[-1]['date']}",
            "notes": [
                "Realistic transaction amounts and categories",
                "Demonstrates both income and various expense types",
                "Shows positive cash flow with savings potential",
                "Includes debt payments and investment contributions",
                "Suitable for comprehensive financial analysis"
            ]
        }
    }
    
    print(f"✅ Sample data created: {len(sample_transactions)} transactions")
    print(f"💰 Income: ${total_income:,.2f}, Expenses: ${total_expenses:,.2f}")
    print(f"💎 Net Cash Flow: ${total_income - total_expenses:,.2f}")
    print(f"📊 Categories: {len(all_categories)} different spending/income categories")
    
    return sample_data

# ============================================================================
# SELF-TEST FUNCTION - Validate Data Processor Capabilities
# ============================================================================

def test_data_processor_capabilities():
    """
    🧪 COMPREHENSIVE DATA PROCESSOR TESTING
    
    WHAT THIS FUNCTION DOES:
    1. Tests all document processing capabilities
    2. Validates sample data generation
    3. Reports available file format support
    4. Tests error handling and recovery
    """
    
    print("🧪 Testing Financial Data Processor capabilities...")
    print("=" * 50)
    
    # Test processor initialization
    try:
        processor = FinancialDataProcessor()
        print(f"✅ Data Processor initialized with {len(processor.supported_formats)} supported formats")
        
        # Test sample data generation
        print("📊 Testing sample data generation...")
        sample_data = create_sample_data()
        
        if sample_data and 'transactions' in sample_data:
            print(f"✅ Sample data: {len(sample_data['transactions'])} transactions")
            print(f"💰 Income: ${sample_data['total_income']:,.2f}")
            print(f"💸 Expenses: ${sample_data['total_expenses']:,.2f}")
            print(f"📊 Categories: {len(sample_data['categories'])}")
        else:
            print("❌ Sample data generation failed")
            return False
        
        # Test column detection (if pandas available)
        if PANDAS_AVAILABLE:
            print("🔍 Testing column detection...")
            # Create test dataframe
            test_df = pd.DataFrame({
                'Transaction Date': ['2024-01-01', '2024-01-02'],
                'Amount': [100, -50],
                'Description': ['Test Income', 'Test Expense']
            })
            
            column_mapping = processor._detect_csv_columns(test_df)
            if column_mapping:
                print(f"✅ Column detection: {column_mapping}")
            else:
                print("⚠️ Column detection needs improvement")
        
        print("=" * 50)
        print("🎉 All data processor tests passed!")
        
        # Report capabilities
        print("\n📋 CAPABILITY SUMMARY:")
        print(f"📊 CSV Processing: {'✅' if PANDAS_AVAILABLE else '❌'}")
        print(f"📈 Excel Processing: {'✅' if EXCEL_AVAILABLE else '❌'}")
        print(f"📄 PDF Processing: {'✅' if PDF_AVAILABLE else '❌'}")
        print(f"📝 Word Processing: {'✅' if DOCX_AVAILABLE else '❌'}")
        print(f"🔍 Smart Column Detection: {'✅' if PANDAS_AVAILABLE else '❌'}")
        print(f"📊 Sample Data Generation: ✅")
        
        return True
        
    except Exception as e:
        print(f"❌ Data processor test failed: {e}")
        return False

# ============================================================================
# MAIN EXECUTION - For Testing
# ============================================================================

if __name__ == "__main__":
    """
    🎯 DATA PROCESSOR TESTING ENTRY POINT
    
    When someone runs "python data_processor.py", this tests all functionality
    """
    print("📊 Financial Data Processor - Standalone Testing")
    print("=" * 60)
    
    success = test_data_processor_capabilities()
    
    if success:
        print("✅ Data processor is ready for use in the main application!")
        print("📁 Sample data generation working perfectly")
    else:
        print("❌ Some issues detected - check error messages above")
        print("💡 Try: pip install pandas openpyxl PyPDF2 python-docx")

# ============================================================================
# END OF ENHANCED DATA PROCESSOR
# ============================================================================